"""
金融制度文本专用分块器
Financial Regulation Text Splitter for LangChain

适配场景：
1. 金融监管文件通常采用"章-条-款-项"的层级结构
2. 需要保持法规条文的完整性，避免语义割裂
3. 过滤过短片段，确保每个chunk包含完整的监管要求
"""

import re
from typing import List, Optional
from langchain_text_splitters import TextSplitter
from langchain_core.documents import Document


class FinancialRegulationSplitter(TextSplitter):
    """
    金融制度文本分块器
    
    核心特性：
    - 按照金融监管文件的层级结构（章/条/款/项）进行智能分割
    - 保留完整的法规条文，避免破坏语义完整性
    - 过滤过短片段，确保知识点的有效性
    """
    
    def __init__(
        self,
        separators: Optional[List[str]] = None,
        min_chunk_size: int = 0,
        keep_separator: bool = True,
        **kwargs
    ):
        """
        初始化分块器
        
        Args:
            separators: 分隔符列表，按优先级排序（默认：章>条>款>项>句号>分号）
            min_chunk_size: 最小分块大小（字符数），设为0表示不过滤
            keep_separator: 是否保留分隔符（保留可维持文档结构）
        """
        super().__init__(**kwargs)
        
        # 金融监管文件的典型层级结构分隔符
        # 优先级：章 > 条 > 款 > 项 > 句子级分隔符
        self.separators = separators or [
            r"[零一二三四五六七八九十百千\d]+[、]",
            r"第[零一二三四五六七八九十百千\d]+条",  # 条款标题
            r"第[零一二三四五六七八九十百千\d]+款",  # 款项标题
            r"[（(][一二三四五六七八九十\d]+[)）]",  # 项（括号形式）
            r"第[零一二三四五六七八九十百千\d]+章",  # 章节标题
            # r"。",  # 句号
            # r"；",  # 分号
        ]
        
        self.min_chunk_size = min_chunk_size
        self.keep_separator = keep_separator
    
    def split_text(self, text: str) -> List[str]:
        """
        分割文本为多个chunk
        
        金融场景适配：
        1. 优先按章节结构分割，保持监管框架完整性
        2. 递归分割过大片段，确保检索效率
        3. 可选过滤过短片段（通过min_chunk_size控制）
        
        Args:
            text: 待分割的文本
            
        Returns:
            分割后的文本片段列表
        """
        chunks = self._split_text_recursive(text, self.separators)
        
        # 过滤过短片段（如果设置了min_chunk_size）
        if self.min_chunk_size > 0:
            filtered_chunks = [
                chunk.strip() 
                for chunk in chunks 
                if len(chunk.strip()) >= self.min_chunk_size
            ]
        else:
            filtered_chunks = [chunk.strip() for chunk in chunks if chunk.strip()]
        
        return filtered_chunks
    
    def _split_text_recursive(
        self, 
        text: str, 
        separators: List[str]
    ) -> List[str]:
        """
        递归分割文本
        
        策略：
        1. 使用当前优先级最高的分隔符进行分割
        2. 如果分割后的片段仍然过大，使用下一级分隔符继续分割
        3. 保留分隔符以维持文档结构的可读性
        """
        if not separators:
            return [text]
        
        separator = separators[0]
        remaining_separators = separators[1:]
        
        # 使用正则表达式查找所有分隔符位置
        pattern = re.compile(separator)
        matches = list(pattern.finditer(text))
        
        if not matches:
            # 当前分隔符未找到，尝试下一级分隔符
            if remaining_separators:
                return self._split_text_recursive(text, remaining_separators)
            return [text]
        
        chunks = []
        last_end = 0
        
        for match in matches:
            # 提取分隔符之前的内容
            before_sep = text[last_end:match.start()].strip()
            if before_sep:
                chunks.append(before_sep)
            
            # 保留分隔符（金融文档中"第X条"等标识符是重要的结构信息）
            if self.keep_separator:
                separator_text = match.group()
                # 查找分隔符后的内容直到下一个分隔符
                next_match_start = matches[matches.index(match) + 1].start() if matches.index(match) + 1 < len(matches) else len(text)
                chunk_with_sep = text[match.start():next_match_start].strip()
                
                # 如果chunk过大，递归分割
                if len(chunk_with_sep) > 300 and remaining_separators:
                    sub_chunks = self._split_text_recursive(chunk_with_sep, remaining_separators)
                    chunks.extend(sub_chunks)
                else:
                    if chunk_with_sep:
                        chunks.append(chunk_with_sep)
                
                last_end = next_match_start
        
        # 处理最后一个分隔符之后的内容
        if last_end < len(text):
            remaining = text[last_end:].strip()
            if remaining:
                chunks.append(remaining)
        
        return chunks


def clean_financial_text(text: str) -> str:
    """
    清洗金融文档文本
    
    金融场景适配：
    1. 去除常见的水印文字（如"内部资料"、"机密"等）
    2. 去除页眉页脚（页码、日期等）
    3. 规范化空白字符
    4. 去除PDF提取时产生的异常字符
    
    Args:
        text: 原始文本
        
    Returns:
        清洗后的文本
    """
    # 去除常见水印文字
    watermark_patterns = [
        r"内部资料",
        r"机密文件",
        r"仅供内部使用",
        r"CONFIDENTIAL",
        r"INTERNAL USE ONLY",
        r"草稿|DRAFT",
    ]
    
    for pattern in watermark_patterns:
        text = re.sub(pattern, "", text, flags=re.IGNORECASE)
    
    # 去除页眉页脚常见模式
    # 1. 页码（如：第1页、Page 1、-1-等）
    text = re.sub(r"第\s*\d+\s*页", "", text)
    text = re.sub(r"Page\s*\d+", "", text, flags=re.IGNORECASE)
    text = re.sub(r"-\s*\d+\s*-", "", text)
    
    # 2. 日期格式（如：2024-01-01、2024年1月1日）
    text = re.sub(r"\d{4}[-年]\d{1,2}[-月]\d{1,2}日?", "", text)
    
    # 3. 常见页眉页脚关键词
    text = re.sub(r"(页眉|页脚|Header|Footer)", "", text, flags=re.IGNORECASE)
    
    # 去除PDF提取时的异常字符
    text = re.sub(r"[\x00-\x08\x0b-\x0c\x0e-\x1f\x7f]", "", text)
    
    # 规范化空白字符
    # 1. 将多个空格替换为单个空格
    text = re.sub(r" +", " ", text)
    
    # 2. 将多个换行符替换为最多两个（保留段落结构）
    text = re.sub(r"\n{3,}", "\n\n", text)
    
    # 3. 去除行首行尾空白
    lines = [line.strip() for line in text.split("\n")]
    text = "\n".join(lines)
    
    # 去除首尾空白
    text = text.strip()
    
    return text


def load_financial_document(
    file_path: str,
    clean_text: bool = True
) -> Document:
    """
    加载金融文档（支持DOCX/PDF/TXT格式）
    
    金融场景适配：
    1. PDF：使用PyMuPDF提取，保留文本结构
    2. DOCX：使用python-docx提取，保留段落格式
    3. TXT：直接读取，适用于已处理的纯文本监管文件
    
    Args:
        file_path: 文件路径
        clean_text: 是否进行文本清洗
        
    Returns:
        LangChain Document对象
    """
    import os
    from pathlib import Path
    
    file_ext = Path(file_path).suffix.lower()
    
    try:
        if file_ext == ".pdf":
            # 使用PyMuPDF (fitz)提取PDF文本
            import fitz
            
            doc = fitz.open(file_path)
            text_parts = []
            
            for page_num in range(len(doc)):
                page = doc[page_num]
                # 提取文本，保留布局
                text = page.get_text("text")
                text_parts.append(text)
            
            doc.close()
            full_text = "\n".join(text_parts)
            
        elif file_ext == ".docx":
            # 使用python-docx提取DOCX文本
            from docx import Document as DocxDocument
            
            doc = DocxDocument(file_path)
            text_parts = []
            
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
            
            full_text = "\n".join(text_parts)
            
        elif file_ext == ".txt":
            # 直接读取TXT文件
            with open(file_path, "r", encoding="utf-8") as f:
                full_text = f.read()
        
        else:
            raise ValueError(f"不支持的文件格式: {file_ext}，仅支持 .pdf, .docx, .txt")
        
        # 文本清洗
        if clean_text:
            full_text = clean_financial_text(full_text)
        
        # 创建LangChain Document对象
        metadata = {
            "source": file_path,
            "file_type": file_ext,
            "file_name": os.path.basename(file_path)
        }
        
        return Document(page_content=full_text, metadata=metadata)
    
    except Exception as e:
        raise Exception(f"加载文档失败: {file_path}, 错误: {str(e)}")
