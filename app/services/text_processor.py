"""
FinRegQA 文本处理服务
Text Processing Service
"""
import re
from typing import List
from langchain_core.documents import Document


class TextSplitterService:
    """金融制度文本分块服务"""
    
    def __init__(self, separators: List[str] = None, min_chunk_size: int = 0, keep_separator: bool = True):
        self.separators = separators or [
            r"第[零一二三四五六七八九十百千\d]+条",
            r"第[零一二三四五六七八九十百千\d]+款",
            r"第[零一二三四五六七八九十百千\d]+章",
        ]
        self.min_chunk_size = min_chunk_size
        self.keep_separator = keep_separator
    
    def split_text(self, text: str) -> List[str]:
        """分割文本"""
        chunks = self._split_text_recursive(text, self.separators)
        if self.min_chunk_size > 0:
            return [chunk.strip() for chunk in chunks if len(chunk.strip()) >= self.min_chunk_size]
        return [chunk.strip() for chunk in chunks if chunk.strip()]
    
    def _split_text_recursive(self, text: str, separators: List[str]) -> List[str]:
        """递归分割文本"""
        if not separators:
            return [text]
        
        separator = separators[0]
        remaining_separators = separators[1:]
        pattern = re.compile(separator)
        matches = list(pattern.finditer(text))
        
        if not matches:
            if remaining_separators:
                return self._split_text_recursive(text, remaining_separators)
            return [text]
        
        chunks = []
        last_end = 0
        
        for match in matches:
            before_sep = text[last_end:match.start()].strip()
            if before_sep:
                chunks.append(before_sep)
            
            if self.keep_separator:
                next_match_start = matches[matches.index(match) + 1].start() if matches.index(match) + 1 < len(matches) else len(text)
                chunk_with_sep = text[match.start():next_match_start].strip()
                
                if len(chunk_with_sep) > 2000 and remaining_separators:
                    chunks.extend(self._split_text_recursive(chunk_with_sep, remaining_separators))
                else:
                    if chunk_with_sep:
                        chunks.append(chunk_with_sep)
                last_end = next_match_start
        
        if last_end < len(text):
            remaining = text[last_end:].strip()
            if remaining:
                chunks.append(remaining)
        
        return chunks


def clean_financial_text(text: str) -> str:
    """清洗金融文档文本"""
    watermark_patterns = [r"内部资料", r"机密文件", r"仅供内部使用", r"CONFIDENTIAL", r"INTERNAL USE ONLY", r"草稿|DRAFT"]
    for pattern in watermark_patterns:
        text = re.sub(pattern, "", text, flags=re.IGNORECASE)
    
    text = re.sub(r"第\s*\d+\s*页", "", text)
    text = re.sub(r"Page\s*\d+", "", text, flags=re.IGNORECASE)
    text = re.sub(r"-\s*\d+\s*-", "", text)
    text = re.sub(r"\d{4}[-年]\d{1,2}[-月]\d{1,2}日?", "", text)
    text = re.sub(r"[\x00-\x08\x0b-\x0c\x0e-\x1f\x7f]", "", text)
    text = re.sub(r" +", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    lines = [line.strip() for line in text.split("\n")]
    text = "\n".join(lines)
    return text.strip()


def load_financial_document(file_path: str, clean_text: bool = True) -> Document:
    """加载金融文档（支持DOCX/PDF/TXT格式）"""
    import os
    from pathlib import Path
    
    file_ext = Path(file_path).suffix.lower()
    
    try:
        if file_ext == ".pdf":
            import fitz
            doc = fitz.open(file_path)
            text_parts = [doc[page_num].get_text("text") for page_num in range(len(doc))]
            doc.close()
            full_text = "\n".join(text_parts)
        elif file_ext == ".docx":
            from docx import Document as DocxDocument
            doc = DocxDocument(file_path)
            text_parts = [para.text for para in doc.paragraphs if para.text.strip()]
            full_text = "\n".join(text_parts)
        elif file_ext == ".txt":
            with open(file_path, "r", encoding="utf-8") as f:
                full_text = f.read()
        else:
            raise ValueError(f"不支持的文件格式: {file_ext}，仅支持 .pdf, .docx, .txt")
        
        if clean_text:
            full_text = clean_financial_text(full_text)
        
        return Document(page_content=full_text, metadata={"source": file_path, "file_type": file_ext, "file_name": os.path.basename(file_path)})
    except Exception as e:
        raise Exception(f"加载文档失败: {file_path}, 错误: {str(e)}")
